from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Query, Request, Response
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import json
import re
import tempfile
import uuid
from bson import ObjectId
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional
from datetime import datetime, timezone

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

app = FastAPI()
api_router = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

from auth import (
    hash_password, verify_password, create_access_token, create_refresh_token,
    get_current_user, seed_users
)

# ─── SEED DATA ───────────────────────────────────────────────────────────────

MODULES_SEED = [
    {"id":"QA101","start_time":"9:00 AM","week":"Week 1","date":"20-Apr","day":"Monday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Quartile","module":"Introductions/Expectations","instructor":"Cristiano Moura","status":"Presented"},
    {"id":"QA102","start_time":"10:10 AM","week":"Week 1","date":"20-Apr","day":"Monday","shift":"AM","length_hrs":0.5,"format":"Live","channel":"Quartile","module":"System & Accesses","instructor":"Cristiano Moura","status":"Presented"},
    {"id":"QA103","start_time":"11:00 AM","week":"Week 1","date":"20-Apr","day":"Monday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Quartile","module":"Life as a CSM","instructor":"Mari Chagas, Pedro Estrada","status":"Presented"},
    {"id":"QA104","start_time":"1:10 PM","week":"Week 1","date":"20-Apr","day":"Monday","shift":"AM","length_hrs":0.5,"format":"Live","channel":"Quartile","module":"Quartile Portal Overview","instructor":"Cristiano Moura, Pedro Estrada","status":"Presented"},
    {"id":"QA105","start_time":"2:50 PM","week":"Week 1","date":"20-Apr","day":"Monday","shift":"PM","length_hrs":0.5,"format":"Live","channel":"Quartile","module":"Quartile Departments & Roles","instructor":"Cristiano Moura","status":"Presented"},
    {"id":"QA106","start_time":"3:30 PM","week":"Week 1","date":"20-Apr","day":"Monday","shift":"PM","length_hrs":1.0,"format":"Live","channel":"Quartile","module":"Ecomm 101","instructor":"Brian","status":"Presented"},
    {"id":"QA107","start_time":"4:40 PM","week":"Week 1","date":"20-Apr","day":"Monday","shift":"PM","length_hrs":1.0,"format":"OnDemand","channel":"AMZ","module":"AMZ Intro: ASIN Structure","instructor":"Ana Braccialli","status":"Presented"},
    {"id":"QA108","start_time":"5:50 PM","week":"Week 1","date":"20-Apr","day":"Monday","shift":"PM","length_hrs":0.5,"format":"OnDemand","channel":"AMZ","module":"AMZ Intro: Buy Box/Retail Readiness","instructor":"Luciana","status":"Presented"},
    {"id":"QA109","start_time":"9:30 AM","week":"Week 1","date":"21-Apr","day":"Tuesday","shift":"AM","length_hrs":1.0,"format":"OnDemand","channel":"AMZ","module":"Branded Terms vs NB","instructor":"Ana Lacroix","status":"Presented"},
    {"id":"QA110","start_time":"10:40 AM","week":"Week 1","date":"21-Apr","day":"Tuesday","shift":"AM","length_hrs":0.5,"format":"OnDemand","channel":"AMZ","module":"AMZ Intro: Sales Attribution","instructor":"Paulo Junqueira","status":"Presented"},
    {"id":"QA111","start_time":"12:20 PM","week":"Week 1","date":"21-Apr","day":"Tuesday","shift":"PM","length_hrs":1.0,"format":"OnDemand","channel":"AMZ","module":"AMZ Intro: Placements (SD, SP, SB)","instructor":"Augusto Senna","status":"Presented"},
    {"id":"QA112","start_time":"1:30 PM","week":"Week 1","date":"21-Apr","day":"Tuesday","shift":"PM","length_hrs":1.0,"format":"OnDemand","channel":"AMZ","module":"AMZ Intro: ACOS/TACOS","instructor":"Ana Braganca","status":"Presented"},
    {"id":"QA113","start_time":"2:40 PM","week":"Week 1","date":"21-Apr","day":"Tuesday","shift":"PM","length_hrs":1.0,"format":"OnDemand","channel":"AMZ","module":"Campaign Behavior","instructor":"Thiago Castro","status":"Presented"},
    {"id":"QA114","start_time":"3:50 PM","week":"Week 1","date":"21-Apr","day":"Tuesday","shift":"PM","length_hrs":1.0,"format":"OnDemand","channel":"AMZ","module":"Naming Convention","instructor":"Thiago Castro","status":"Presented"},
    {"id":"QA115","start_time":"9:30 AM","week":"Week 1","date":"22-Apr","day":"Wednesday","shift":"AM","length_hrs":0.5,"format":"Live","channel":"AMZ","module":"AMZ Intro: Vendor vs Seller","instructor":"Tiago Fronza","status":"Presented"},
    {"id":"QA116","start_time":"10:10 AM","week":"Week 1","date":"22-Apr","day":"Wednesday","shift":"AM","length_hrs":1.0,"format":"OnDemand","channel":"AMZ","module":"Sponsored Products 101","instructor":"Felipe Tahara","status":"Presented"},
    {"id":"QA117","start_time":"12:20 PM","week":"Week 1","date":"22-Apr","day":"Wednesday","shift":"PM","length_hrs":3.0,"format":"Live","channel":"AMZ","module":"Sponsored Products - Advanced","instructor":"Felipe Tahara","status":"Presented"},
    {"id":"QA118","start_time":"3:30 PM","week":"Week 1","date":"22-Apr","day":"Wednesday","shift":"PM","length_hrs":0.5,"format":"OnDemand","channel":"Quartile","module":"ICP","instructor":"Fernando Gamba","status":"Presented"},
    {"id":"QA119","start_time":"9:30 AM","week":"Week 1","date":"23-Apr","day":"Thursday","shift":"AM","length_hrs":1.0,"format":"OnDemand","channel":"AMZ","module":"Sponsored Brands 101","instructor":"Frederico Cappellato","status":"Presented"},
    {"id":"QA120","start_time":"11:40 AM","week":"Week 1","date":"23-Apr","day":"Thursday","shift":"PM","length_hrs":3.0,"format":"Live","channel":"AMZ","module":"Sponsored Brands Advanced","instructor":"Frederico Cappellato","status":"Presented"},
    {"id":"QA121","start_time":"9:30 AM","week":"Week 1","date":"24-Apr","day":"Friday","shift":"AM","length_hrs":1.5,"format":"OnDemand","channel":"AMZ","module":"DSP/SD 101","instructor":"Regina Reker","status":"Presented"},
    {"id":"QA122","start_time":"12:50 PM","week":"Week 1","date":"24-Apr","day":"Friday","shift":"PM","length_hrs":1.0,"format":"Live","channel":"Quartile","module":"Buddy Introduction","instructor":"Buddy","status":"Presented"},
    {"id":"QA123","start_time":"1:30 PM","week":"Week 1","date":"24-Apr","day":"Friday","shift":"PM","length_hrs":0.5,"format":"Live","channel":"Quartile","module":"Leadership Fireside Chat","instructor":"Amanda, Phil, Roberto","status":"Presented"},
    {"id":"QA124","start_time":"2:00 PM","week":"Week 1","date":"24-Apr","day":"Friday","shift":"PM","length_hrs":3.0,"format":"Live","channel":"AMZ","module":"Assessment 1: Channels/Placements","instructor":"Brenda Mentz","status":"Presented"},
    {"id":"QA201","start_time":"9:30 AM","week":"Week 2","date":"26-Apr","day":"Monday","shift":"AM","length_hrs":2.0,"format":"Live","channel":"Portal","module":"Quartile Portal I - ACOS Hierarchy","instructor":"Thomas Shimidt","status":"Presented"},
    {"id":"QA202","start_time":"12:40 PM","week":"Week 2","date":"26-Apr","day":"Monday","shift":"PM","length_hrs":2.0,"format":"Live","channel":"Portal","module":"Quartile Portal II - Campaign creation","instructor":"Thomas Shimidt","status":"Presented"},
    {"id":"QA203","start_time":"9:30 AM","week":"Week 2","date":"27-Apr","day":"Tuesday","shift":"AM","length_hrs":2.0,"format":"Live","channel":"Portal","module":"Quartile Portal III - Imported campaigns","instructor":"Thomas Shimidt","status":"Presented"},
    {"id":"QA204","start_time":"12:40 PM","week":"Week 2","date":"27-Apr","day":"Tuesday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"AMZ","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA205","start_time":"9:30 AM","week":"Week 2","date":"28-Apr","day":"Wednesday","shift":"AM","length_hrs":2.0,"format":"Live","channel":"Portal","module":"Quartile Portal IV - Reports","instructor":"Guilherme Decimo","status":"Scheduled"},
    {"id":"QA206","start_time":"12:40 PM","week":"Week 2","date":"28-Apr","day":"Wednesday","shift":"PM","length_hrs":2.0,"format":"Live","channel":"Quartile","module":"Portal Q&A and Practical Exercises","instructor":"Luis Saicali","status":"Scheduled"},
    {"id":"QA207","start_time":"2:50 PM","week":"Week 2","date":"28-Apr","day":"Wednesday","shift":"PM","length_hrs":1.0,"format":"Live","channel":"Soft Skills","module":"Objection Handling","instructor":"Mari Chagas","status":"Scheduled"},
    {"id":"QA208","start_time":"4:00 PM","week":"Week 2","date":"28-Apr","day":"Wednesday","shift":"PM","length_hrs":1.0,"format":"Live","channel":"Soft Skills","module":"Negotiation","instructor":"Mari Chagas","status":"Scheduled"},
    {"id":"QA209","start_time":"9:30 AM","week":"Week 2","date":"29-Apr","day":"Thursday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"AMZ","module":"DSP strategist","instructor":"Regina Reker","status":"Scheduled"},
    {"id":"QA210","start_time":"10:40 AM","week":"Week 2","date":"29-Apr","day":"Thursday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"AMZ","module":"DSP AMC Audiences + Reports","instructor":"Victor Anger","status":"Scheduled"},
    {"id":"QA211","start_time":"9:30 AM","week":"Week 2","date":"30-Apr","day":"Friday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"AMZ","module":"DSP Situations & Funnel","instructor":"Victor Anger","status":"Scheduled"},
    {"id":"QA212","start_time":"11:40 AM","week":"Week 2","date":"30-Apr","day":"Friday","shift":"PM","length_hrs":0.5,"format":"Live","channel":"Quartile","module":"Leadership Fireside Chat","instructor":"Chris Corbet","status":"Scheduled"},
    {"id":"QA213","start_time":"12:20 PM","week":"Week 2","date":"30-Apr","day":"Friday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Quartile","module":"Assessment 2: Campaign Mgt","instructor":"Thomas Shimidt","status":"Scheduled"},
    {"id":"QA301","start_time":"9:30 AM","week":"Week 3","date":"4-May","day":"Monday","shift":"AM","length_hrs":1.0,"format":"OnDemand","channel":"Google","module":"Direct To Consumer","instructor":"Carlos Braz","status":"Scheduled"},
    {"id":"QA302","start_time":"10:40 AM","week":"Week 3","date":"4-May","day":"Monday","shift":"AM","length_hrs":0.5,"format":"OnDemand","channel":"Google","module":"Merchant Center Overview","instructor":"Bruno Matumoto","status":"Scheduled"},
    {"id":"QA303","start_time":"11:20 AM","week":"Week 3","date":"4-May","day":"Monday","shift":"AM","length_hrs":0.5,"format":"Live","channel":"Google Ads","module":"Feed Overview","instructor":"Frank Savena","status":"Scheduled"},
    {"id":"QA304","start_time":"1:00 PM","week":"Week 3","date":"4-May","day":"Monday","shift":"PM","length_hrs":0.5,"format":"Live","channel":"Google","module":"GA4 Overview + Insights","instructor":"Pat Bradley","status":"Scheduled"},
    {"id":"QA305","start_time":"1:40 PM","week":"Week 3","date":"4-May","day":"Monday","shift":"PM","length_hrs":1.0,"format":"OnDemand","channel":"Google Ads","module":"Google Ads Overview","instructor":"Luisa Conti","status":"Scheduled"},
    {"id":"QA306","start_time":"2:50 PM","week":"Week 3","date":"4-May","day":"Monday","shift":"PM","length_hrs":1.0,"format":"Live","channel":"Google Ads","module":"Paid Search","instructor":"Jose Adorno","status":"Scheduled"},
    {"id":"QA307","start_time":"9:30 AM","week":"Week 3","date":"5-May","day":"Tuesday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Google Ads","module":"Standard Shopping","instructor":"Rebecca Rush","status":"Scheduled"},
    {"id":"QA308","start_time":"10:40 AM","week":"Week 3","date":"5-May","day":"Tuesday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Google Ads","module":"PMAX","instructor":"Rebecca Rush","status":"Scheduled"},
    {"id":"QA309","start_time":"12:50 PM","week":"Week 3","date":"5-May","day":"Tuesday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Quartile","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA310","start_time":"9:30 AM","week":"Week 3","date":"6-May","day":"Wednesday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Google Ads","module":"PLA Segmentation Strategy","instructor":"Jake Damico","status":"Scheduled"},
    {"id":"QA311","start_time":"10:40 AM","week":"Week 3","date":"6-May","day":"Wednesday","shift":"AM","length_hrs":0.5,"format":"Live","channel":"Google Ads","module":"Quartile Tech (DTC)","instructor":"Sabrina Morais","status":"Scheduled"},
    {"id":"QA312","start_time":"12:20 PM","week":"Week 3","date":"6-May","day":"Wednesday","shift":"PM","length_hrs":1.0,"format":"Live","channel":"Google Ads","module":"Dashboards","instructor":"Cristiano Moura","status":"Scheduled"},
    {"id":"QA313","start_time":"9:30 AM","week":"Week 3","date":"7-May","day":"Thursday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Bing","module":"Bing Overview","instructor":"Carlos Braz","status":"Scheduled"},
    {"id":"QA314","start_time":"10:40 AM","week":"Week 3","date":"7-May","day":"Thursday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Meta","module":"Meta Overview","instructor":"Luca Dias","status":"Scheduled"},
    {"id":"QA315","start_time":"12:50 PM","week":"Week 3","date":"7-May","day":"Thursday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Quartile","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA316","start_time":"9:30 AM","week":"Week 3","date":"8-May","day":"Friday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"DTC","module":"Health Checks","instructor":"Julia Baldo","status":"Scheduled"},
    {"id":"QA317","start_time":"10:40 AM","week":"Week 3","date":"8-May","day":"Friday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"DTC","module":"Performance Calls","instructor":"Kerianne Kistner","status":"Scheduled"},
    {"id":"QA318","start_time":"12:50 PM","week":"Week 3","date":"8-May","day":"Friday","shift":"PM","length_hrs":0.5,"format":"Live","channel":"Quartile","module":"Leadership Fireside Chat","instructor":"Laurence","status":"Scheduled"},
    {"id":"QA319","start_time":"1:30 PM","week":"Week 3","date":"8-May","day":"Friday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Google","module":"Assessment 3: Google Analysis","instructor":"Pat Bradley","status":"Scheduled"},
    {"id":"QA401","start_time":"9:30 AM","week":"Week 4","date":"11-May","day":"Monday","shift":"AM","length_hrs":2.0,"format":"Live","channel":"Walmart","module":"Walmart Overview","instructor":"Luciano Ferrareze, Ana Maiczuk","status":"Scheduled"},
    {"id":"QA402","start_time":"12:40 PM","week":"Week 4","date":"11-May","day":"Monday","shift":"PM","length_hrs":0.5,"format":"OnDemand","channel":"Tech","module":"Pro Suite Overview","instructor":"Felipe Tahara","status":"Scheduled"},
    {"id":"QA403","start_time":"1:20 PM","week":"Week 4","date":"11-May","day":"Monday","shift":"PM","length_hrs":1.0,"format":"Live","channel":"TikTok","module":"TikTok 101","instructor":"Luca Dias","status":"Scheduled"},
    {"id":"QA404","start_time":"9:30 AM","week":"Week 4","date":"12-May","day":"Tuesday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Soft Skills","module":"Upsell Process, CG","instructor":"Joao Victor, Julia Fortino","status":"Scheduled"},
    {"id":"QA405","start_time":"10:40 AM","week":"Week 4","date":"12-May","day":"Tuesday","shift":"AM","length_hrs":1.0,"format":"OnDemand","channel":"Tech","module":"Promo Management Overview","instructor":"Matheus Cardoso","status":"Scheduled"},
    {"id":"QA406","start_time":"12:50 PM","week":"Week 4","date":"12-May","day":"Tuesday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Quartile","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA407","start_time":"9:30 AM","week":"Week 4","date":"13-May","day":"Wednesday","shift":"AM","length_hrs":1.0,"format":"OnDemand","channel":"Sciene","module":"Sciene 101","instructor":"Ricardo Morandini","status":"Scheduled"},
    {"id":"QA408","start_time":"10:40 AM","week":"Week 4","date":"13-May","day":"Wednesday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"ChatGPT","module":"AI Tools in Practice","instructor":"Camila Malacarne","status":"Scheduled"},
    {"id":"QA409","start_time":"12:50 PM","week":"Week 4","date":"13-May","day":"Wednesday","shift":"PM","length_hrs":2.0,"format":"Live","channel":"Sciene","module":"Sciene Companion in Practice","instructor":"Camila Malacarne","status":"Scheduled"},
    {"id":"QA410","start_time":"3:00 PM","week":"Week 4","date":"13-May","day":"Wednesday","shift":"PM","length_hrs":1.0,"format":"Live","channel":"Quartile","module":"Customer Journey at Quartile","instructor":"Joao Victor, Julia Fortino","status":"Scheduled"},
    {"id":"QA411","start_time":"9:30 AM","week":"Week 4","date":"14-May","day":"Thursday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Quartile","module":"Account Transition Process","instructor":"Cristiano Moura","status":"Scheduled"},
    {"id":"QA412","start_time":"11:40 AM","week":"Week 4","date":"14-May","day":"Thursday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Quartile","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA413","start_time":"9:30 AM","week":"Week 4","date":"15-May","day":"Friday","shift":"AM","length_hrs":2.0,"format":"Live","channel":"Claude","module":"Assessment 4: AI Tools","instructor":"Camila Malacarne","status":"Scheduled"},
    {"id":"QA414","start_time":"12:40 PM","week":"Week 4","date":"15-May","day":"Friday","shift":"PM","length_hrs":0.5,"format":"Live","channel":"Quartile","module":"Leadership Fireside Chat","instructor":"Sylvio","status":"Scheduled"},
    {"id":"QA501","start_time":"9:30 AM","week":"Week 5","date":"16-May","day":"Monday","shift":"AM","length_hrs":1.0,"format":"OnDemand","channel":"Excel","module":"Excel Advanced","instructor":"Alberto Romero","status":"Scheduled"},
    {"id":"QA502","start_time":"10:40 AM","week":"Week 5","date":"16-May","day":"Monday","shift":"AM","length_hrs":1.0,"format":"OnDemand","channel":"AMZ","module":"Account Analysis 101 (AMZ)","instructor":"Luis Saicali","status":"Scheduled"},
    {"id":"QA503","start_time":"12:50 PM","week":"Week 5","date":"16-May","day":"Monday","shift":"PM","length_hrs":3.0,"format":"Live","channel":"AMZ","module":"Account Analysis Advanced (AMZ)","instructor":"Iago Gomes","status":"Scheduled"},
    {"id":"QA504","start_time":"4:00 PM","week":"Week 5","date":"16-May","day":"Monday","shift":"PM","length_hrs":1.0,"format":"OnDemand","channel":"DTC","module":"Account Analysis 101 (DTC)","instructor":"Emily Foster","status":"Scheduled"},
    {"id":"QA505","start_time":"9:30 AM","week":"Week 5","date":"17-May","day":"Tuesday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"DTC","module":"Account Analysis Advanced (DTC)","instructor":"Pat Bradley","status":"Scheduled"},
    {"id":"QA506","start_time":"10:40 AM","week":"Week 5","date":"17-May","day":"Tuesday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Google","module":"Quarterly Reviews","instructor":"Cristiano Moura","status":"Scheduled"},
    {"id":"QA507","start_time":"12:50 PM","week":"Week 5","date":"17-May","day":"Tuesday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Quartile","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA508","start_time":"9:30 AM","week":"Week 5","date":"18-May","day":"Wednesday","shift":"AM","length_hrs":2.0,"format":"Live","channel":"Soft Skills","module":"Data into Insights","instructor":"Felipe Tahara","status":"Scheduled"},
    {"id":"QA509","start_time":"11:40 AM","week":"Week 5","date":"18-May","day":"Wednesday","shift":"AM","length_hrs":0.5,"format":"OnDemand","channel":"AMZ","module":"AMC 101","instructor":"Mickaela Pulcherio","status":"Scheduled"},
    {"id":"QA510","start_time":"1:20 PM","week":"Week 5","date":"18-May","day":"Wednesday","shift":"PM","length_hrs":2.0,"format":"Live","channel":"Soft Skills","module":"Measurement Storytelling","instructor":"Amanda Gomez","status":"Scheduled"},
    {"id":"QA511","start_time":"3:30 PM","week":"Week 5","date":"18-May","day":"Wednesday","shift":"PM","length_hrs":2.0,"format":"Assessment","channel":"Google","module":"Assessment: Account Analysis Task","instructor":"Pat Bradley","status":"Scheduled"},
    {"id":"QA512","start_time":"5:40 PM","week":"Week 5","date":"18-May","day":"Wednesday","shift":"PM","length_hrs":2.0,"format":"Assessment","channel":"AMZ","module":"Assessment: Account Analysis Task","instructor":"Iago Gomes","status":"Scheduled"},
    {"id":"QA513","start_time":"9:30 AM","week":"Week 5","date":"19-May","day":"Thursday","shift":"AM","length_hrs":2.0,"format":"Assessment","channel":"Multichannel","module":"Assessment: Account Analysis Task","instructor":"Iago Gomes","status":"Scheduled"},
    {"id":"QA514","start_time":"11:40 AM","week":"Week 5","date":"19-May","day":"Thursday","shift":"AM","length_hrs":0.5,"format":"OnDemand","channel":"Quartile","module":"Market Intelligence 101","instructor":"Suriane Silva","status":"Scheduled"},
    {"id":"QA515","start_time":"1:20 PM","week":"Week 5","date":"19-May","day":"Thursday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Quartile","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA516","start_time":"9:30 AM","week":"Week 5","date":"20-May","day":"Friday","shift":"AM","length_hrs":0.5,"format":"OnDemand","channel":"Quartile","module":"Forecast 101","instructor":"Vitor Caires","status":"Scheduled"},
    {"id":"QA517","start_time":"10:10 AM","week":"Week 5","date":"20-May","day":"Friday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"AMZ","module":"(Buddy) Prep Role Play: Performance Call","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA518","start_time":"12:20 PM","week":"Week 5","date":"20-May","day":"Friday","shift":"PM","length_hrs":0.5,"format":"Live","channel":"Quartile","module":"Leadership Fireside Chat","instructor":"Ed","status":"Scheduled"},
    {"id":"QA519","start_time":"1:00 PM","week":"Week 5","date":"20-May","day":"Friday","shift":"PM","length_hrs":4.0,"format":"Assessment","channel":"AMZ","module":"Prep Role Play: Performance Call","instructor":"Self","status":"Scheduled"},
    {"id":"QA601","start_time":"9:30 AM","week":"Week 6","date":"26-May","day":"Tuesday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Quartile","module":"Market Intelligence Advanced","instructor":"Suriane Silva","status":"Scheduled"},
    {"id":"QA091","start_time":"11:40 AM","week":"Week 6","date":"26-May","day":"Tuesday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"AMZ","module":"Role Plays: Performance Call","instructor":"Directors TBD","status":"Scheduled"},
    {"id":"QA092","start_time":"9:30 AM","week":"Week 6","date":"27-May","day":"Wednesday","shift":"AM","length_hrs":1.0,"format":"OnDemand","channel":"Quartile","module":"GGS & partners","instructor":"TBD","status":"Scheduled"},
    {"id":"QA093","start_time":"11:40 AM","week":"Week 6","date":"27-May","day":"Wednesday","shift":"PM","length_hrs":2.0,"format":"Live","channel":"AMZ","module":"Role Plays: Performance Call","instructor":"Directors TBD","status":"Scheduled"},
    {"id":"QA094","start_time":"9:30 AM","week":"Week 6","date":"28-May","day":"Thursday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Quartile","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA095","start_time":"9:30 AM","week":"Week 6","date":"29-May","day":"Friday","shift":"AM","length_hrs":1.0,"format":"Live","channel":"Multichannel","module":"(Buddy) QBR Task/Prep Role Play","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA096","start_time":"11:40 AM","week":"Week 6","date":"29-May","day":"Friday","shift":"PM","length_hrs":4.0,"format":"Assessment","channel":"Multichannel","module":"QBR Task/Prep Role Play","instructor":"Self","status":"Scheduled"},
    {"id":"QA701","start_time":"9:30 AM","week":"Week 7","date":"1-Jun","day":"Monday","shift":"AM","length_hrs":2.0,"format":"Live","channel":"Quartile","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA702","start_time":"12:40 PM","week":"Week 7","date":"1-Jun","day":"Monday","shift":"PM","length_hrs":4.0,"format":"Live","channel":"Quartile","module":"Pod / Buddy / Shadow calls","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA703","start_time":"9:30 AM","week":"Week 7","date":"2-Jun","day":"Tuesday","shift":"AM","length_hrs":2.0,"format":"Live","channel":"Multichannel","module":"Review QBR w/Buddy","instructor":"Buddy","status":"Scheduled"},
    {"id":"QA704","start_time":"11:40 AM","week":"Week 7","date":"2-Jun","day":"Tuesday","shift":"AM","length_hrs":4.0,"format":"Live","channel":"MC","module":"Role Plays: QBR","instructor":"Directors TBD","status":"Scheduled"},
    {"id":"QA705","start_time":"9:30 AM","week":"Week 7","date":"3-Jun","day":"Wednesday","shift":"AM","length_hrs":2.0,"format":"Live","channel":"MC","module":"Role Plays: QBR","instructor":"Directors TBD","status":"Scheduled"},
    {"id":"QA706","start_time":"9:30 AM","week":"Week 7","date":"5-Jun","day":"Friday","shift":"PM","length_hrs":0.5,"format":"Live","channel":"Quartile","module":"Graduation: Leadership Fireside Chat","instructor":"DK","status":"Scheduled"},
    {"id":"QA707","start_time":"10:40 AM","week":"Week 7","date":"5-Jun","day":"Friday","shift":"AM","length_hrs":8.0,"format":"On-Site","channel":"Quartile","module":"Graduation","instructor":"Cristiano Moura, Monique","status":"Scheduled"},
]

async def seed_data():
    count = await db.modules.count_documents({})
    if count == 0:
        logger.info("Seeding module data...")
        await db.modules.insert_many(MODULES_SEED)
        logger.info(f"Seeded {len(MODULES_SEED)} modules")

        # Derive instructors from modules
        instructor_map = {}
        for m in MODULES_SEED:
            names = [n.strip() for n in m["instructor"].split(",")]
            for name in names:
                if name not in instructor_map and name not in ["Self", "TBD", "Directors TBD", "Buddy"]:
                    instructor_map[name] = {
                        "name": name,
                        "channels": set(),
                        "modules": [],
                        "total_hrs": 0.0
                    }
                if name in instructor_map:
                    instructor_map[name]["channels"].add(m["channel"])
                    instructor_map[name]["modules"].append(m["module"])
                    instructor_map[name]["total_hrs"] += m["length_hrs"]

        instructors = []
        for name, data in instructor_map.items():
            instructors.append({
                "name": data["name"],
                "channels": list(data["channels"]),
                "modules": list(set(data["modules"])),
                "total_hrs": round(data["total_hrs"], 1),
                "role": "COE Specialist"
            })
        if instructors:
            await db.instructors.insert_many(instructors)
            logger.info(f"Seeded {len(instructors)} instructors")

        # Seed students with enriched KPI schema
        base_students = [
            {"name":"Ana Reyes","role":"Advanced Analytics Lead","country":"Brazil","student_id":"ARCH-9234"},
            {"name":"Marcus Liu","role":"Strategic Operations","country":"USA","student_id":"ARCH-9235"},
            {"name":"Priya Shankar","role":"Marketing Science Analyst","country":"India","student_id":"ARCH-9236"},
            {"name":"Jordan Cole","role":"Campaign Associate","country":"USA","student_id":"ARCH-9237"},
            {"name":"Sofia Martinez","role":"Data Analytics Lead","country":"Spain","student_id":"ARCH-9238"},
            {"name":"Carlos Mendes","role":"Performance Manager","country":"Colombia","student_id":"ARCH-9239"},
        ]
        for s in base_students:
            s.update({
                "academic_progress": 0,
                "sessions": 0,
                "total_hours": 0.0,
                "overall_score": 0.0,
                "modules_attended": [],       # [{"module_id":"QA117","module":"name","date":"22-Apr","score":8,"channel":"AMZ","hrs":3.0}]
                "channel_hours": {},           # {"AMZ":5.0, "Google":3.0, ...}
                "shadow_calls": 0,
                "buddy_tasks": 0,
                "assessment_scores": [],       # [{"module_id":"QA124","title":"Assessment 1","score":85}]
                "skills": {
                    "objection_handling": 0,
                    "negotiation": 0,
                    "data_analysis": 0,
                    "communication": 0,
                    "presentation": 0,
                    "analytical_thinking": 0,
                    "campaign_management": 0,
                    "client_management": 0,
                },
                "strengths_summary": "",
                "improvement_areas": "",
            })
        await db.students.insert_many(base_students)
        logger.info(f"Seeded {len(base_students)} students")

@app.on_event("startup")
async def startup():
    await seed_data()
    await seed_users(db, logger)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

# ─── AUTH ENDPOINTS ──────────────────────────────────────────────────────

class LoginBody(BaseModel):
    email: str
    password: str

@api_router.post("/auth/login")
async def login(body: LoginBody, response: Response):
    user = await db.users.find_one({"email": body.email})
    if not user or not verify_password(body.password, user.get("password_hash", "")):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    uid = str(user["_id"])
    access = create_access_token(uid, user["email"], user["role"])
    refresh = create_refresh_token(uid)
    resp = JSONResponse({"token": access, "user": {"email": user["email"], "name": user["name"], "role": user["role"]}})
    resp.set_cookie("access_token", access, httponly=True, samesite="none", secure=True, max_age=86400)
    resp.set_cookie("refresh_token", refresh, httponly=True, samesite="none", secure=True, max_age=604800)
    return resp

@api_router.post("/auth/logout")
async def logout(response: Response):
    resp = JSONResponse({"message": "Logged out"})
    resp.delete_cookie("access_token")
    resp.delete_cookie("refresh_token")
    return resp

@api_router.get("/auth/me")
async def get_me(request: Request):
    user = await get_current_user(request, db)
    return {"email": user["email"], "name": user["name"], "role": user["role"]}

# ─── ADMIN CRUD ENDPOINTS ────────────────────────────────────────────────

# Helper to check admin role
async def require_admin(request: Request):
    user = await get_current_user(request, db)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return user

# ADMIN: Module CRUD
class ModuleBody(BaseModel):
    id: str
    start_time: str = ""
    week: str = ""
    date: str = ""
    day: str = ""
    shift: str = ""
    length_hrs: float = 0
    format: str = ""
    channel: str = ""
    module: str = ""
    instructor: str = ""
    status: str = "Scheduled"

@api_router.post("/admin/modules")
async def admin_create_module(body: ModuleBody, request: Request):
    await require_admin(request)
    existing = await db.modules.find_one({"id": body.id})
    if existing:
        raise HTTPException(status_code=400, detail=f"Module {body.id} already exists")
    data = body.dict()
    data["created_at"] = datetime.now(timezone.utc).isoformat()
    await db.modules.insert_one(data)
    return {"status": "created", "id": body.id}

@api_router.put("/admin/modules/{module_id}")
async def admin_update_module(module_id: str, body: ModuleBody, request: Request):
    await require_admin(request)
    existing = await db.modules.find_one({"id": module_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Module not found")
    # Save version history
    await db.module_versions.insert_one({"module_id": module_id, "data": {k: v for k, v in existing.items() if k != "_id"}, "saved_at": datetime.now(timezone.utc).isoformat()})
    data = body.dict()
    await db.modules.update_one({"id": module_id}, {"$set": data})
    return {"status": "updated", "id": module_id}

@api_router.delete("/admin/modules/{module_id}")
async def admin_delete_module(module_id: str, request: Request):
    await require_admin(request)
    result = await db.modules.delete_one({"id": module_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Module not found")
    return {"status": "deleted", "id": module_id}

@api_router.get("/admin/modules/{module_id}/versions")
async def admin_module_versions(module_id: str, request: Request):
    await require_admin(request)
    versions = await db.module_versions.find({"module_id": module_id}, {"_id": 0}).sort("saved_at", -1).to_list(50)
    return versions

@api_router.post("/admin/modules/{module_id}/restore/{version_idx}")
async def admin_restore_module(module_id: str, version_idx: int, request: Request):
    await require_admin(request)
    versions = await db.module_versions.find({"module_id": module_id}).sort("saved_at", -1).to_list(50)
    if version_idx >= len(versions):
        raise HTTPException(status_code=404, detail="Version not found")
    old_data = versions[version_idx]["data"]
    await db.modules.update_one({"id": module_id}, {"$set": old_data})
    return {"status": "restored", "id": module_id}

# ADMIN: Student CRUD
class StudentBody(BaseModel):
    name: str
    role: str = ""
    country: str = ""
    student_id: str = ""

@api_router.post("/admin/students")
async def admin_create_student(body: StudentBody, request: Request):
    await require_admin(request)
    data = body.dict()
    data.update({"academic_progress": 0, "sessions": 0, "total_hours": 0.0, "overall_score": 0.0,
        "modules_attended": [], "channel_hours": {}, "shadow_calls": 0, "buddy_tasks": 0,
        "assessment_scores": [], "skills": {"objection_handling":0,"negotiation":0,"data_analysis":0,"communication":0,"presentation":0,"analytical_thinking":0,"campaign_management":0,"client_management":0},
        "strengths_summary": "", "improvement_areas": "", "created_at": datetime.now(timezone.utc).isoformat()})
    await db.students.insert_one(data)
    return {"status": "created", "student_id": body.student_id}

@api_router.put("/admin/students/{student_id}")
async def admin_update_student(student_id: str, body: StudentBody, request: Request):
    await require_admin(request)
    existing = await db.students.find_one({"student_id": student_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Student not found")
    await db.student_versions.insert_one({"student_id": student_id, "data": {k: v for k, v in existing.items() if k != "_id"}, "saved_at": datetime.now(timezone.utc).isoformat()})
    await db.students.update_one({"student_id": student_id}, {"$set": body.dict()})
    return {"status": "updated", "student_id": student_id}

@api_router.delete("/admin/students/{student_id}")
async def admin_delete_student(student_id: str, request: Request):
    await require_admin(request)
    result = await db.students.delete_one({"student_id": student_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Student not found")
    return {"status": "deleted", "student_id": student_id}

@api_router.get("/admin/students/{student_id}/versions")
async def admin_student_versions(student_id: str, request: Request):
    await require_admin(request)
    versions = await db.student_versions.find({"student_id": student_id}, {"_id": 0}).sort("saved_at", -1).to_list(50)
    return versions

@api_router.post("/admin/students/{student_id}/restore/{version_idx}")
async def admin_restore_student(student_id: str, version_idx: int, request: Request):
    await require_admin(request)
    versions = await db.student_versions.find({"student_id": student_id}).sort("saved_at", -1).to_list(50)
    if version_idx >= len(versions):
        raise HTTPException(status_code=404, detail="Version not found")
    await db.students.update_one({"student_id": student_id}, {"$set": versions[version_idx]["data"]})
    return {"status": "restored"}

# ADMIN: Instructor CRUD
class InstructorBody(BaseModel):
    name: str
    channels: list = []
    modules: list = []
    total_hrs: float = 0
    role: str = "COE Specialist"

@api_router.post("/admin/instructors")
async def admin_create_instructor(body: InstructorBody, request: Request):
    await require_admin(request)
    data = body.dict()
    data["created_at"] = datetime.now(timezone.utc).isoformat()
    await db.instructors.insert_one(data)
    return {"status": "created", "name": body.name}

@api_router.put("/admin/instructors/{name}")
async def admin_update_instructor(name: str, body: InstructorBody, request: Request):
    await require_admin(request)
    existing = await db.instructors.find_one({"name": name})
    if not existing:
        raise HTTPException(status_code=404, detail="Instructor not found")
    await db.instructor_versions.insert_one({"name": name, "data": {k: v for k, v in existing.items() if k != "_id"}, "saved_at": datetime.now(timezone.utc).isoformat()})
    await db.instructors.update_one({"name": name}, {"$set": body.dict()})
    return {"status": "updated"}

@api_router.delete("/admin/instructors/{name}")
async def admin_delete_instructor(name: str, request: Request):
    await require_admin(request)
    result = await db.instructors.delete_one({"name": name})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Instructor not found")
    return {"status": "deleted"}

# ADMIN: User/Password Management
@api_router.get("/admin/users")
async def admin_list_users(request: Request):
    await require_admin(request)
    users = await db.users.find({}, {"password_hash": 0}).to_list(100)
    for u in users:
        u["_id"] = str(u["_id"])
    return users

class PasswordResetBody(BaseModel):
    email: str
    new_password: str

@api_router.post("/admin/users/reset-password")
async def admin_reset_password(body: PasswordResetBody, request: Request):
    await require_admin(request)
    user = await db.users.find_one({"email": body.email})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    await db.users.update_one({"email": body.email}, {"$set": {"password_hash": hash_password(body.new_password)}})
    return {"status": "password_reset", "email": body.email}

class CreateUserBody(BaseModel):
    email: str
    password: str
    name: str
    role: str = "student"

@api_router.post("/admin/users")
async def admin_create_user(body: CreateUserBody, request: Request):
    await require_admin(request)
    existing = await db.users.find_one({"email": body.email})
    if existing:
        raise HTTPException(status_code=400, detail="User already exists")
    await db.users.insert_one({
        "email": body.email, "password_hash": hash_password(body.password),
        "name": body.name, "role": body.role, "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"status": "created", "email": body.email}

# ─── FILE READING HELPER ─────────────────────────────────────────────────

async def read_uploaded_file(file: UploadFile) -> str:
    """Read text from .docx or .txt uploads."""
    fname = (file.filename or "").lower()
    if not (fname.endswith(".docx") or fname.endswith(".txt")):
        raise HTTPException(status_code=400, detail="Only .docx and .txt files are supported")
    content = await file.read()
    if fname.endswith(".txt"):
        return content.decode("utf-8", errors="replace")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    try:
        import docx
        doc = docx.Document(tmp_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    finally:
        os.unlink(tmp_path)

# ─── ASSESSMENT SCORE ENDPOINT ───────────────────────────────────────────

@api_router.post("/modules/{module_id}/assessment")
async def deliver_assessment(module_id: str, file: UploadFile = File(...), request: Request = None):
    module = await db.modules.find_one({"id": module_id}, {"_id": 0})
    if not module:
        raise HTTPException(status_code=404, detail="Module not found")

    text = await read_uploaded_file(file)

    try:
        from openai import AsyncOpenAI as _OpenAI
        api_key = os.environ.get("EMERGENT_LLM_KEY")
        system_prompt = """You are an assessment evaluator for Quartile Academy. Grade the student's assessment submission. Return ONLY valid JSON:
{"student_name":"Name","overall_score":85,"knowledge":{"score":8,"feedback":"detail"},"data_clarity":{"score":7,"feedback":"detail"},"communication":{"score":9,"feedback":"detail"},"soft_skills":{"score":8,"feedback":"detail"},"critical_thinking":{"score":7,"feedback":"detail"},"problem_solving":{"score":8,"feedback":"detail"},"technical_accuracy":{"score":9,"feedback":"detail"},"time_management":{"score":7,"feedback":"detail"},"strategic_thinking":{"score":8,"feedback":"detail"},"overall_feedback":"2-3 sentence summary","strengths":["s1","s2"],"areas_for_improvement":["a1","a2"]}
Score each skill 1-10. Be thorough and fair."""

        _client = _OpenAI(api_key=api_key)
        user_text = f"Assessment for {module.get('module', '')} ({module.get('channel', '')}):\n\n{text[:15000]}"
        _completion = await _client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_text}
            ]
        )
        resp = _completion.choices[0].message.content
        resp_text = resp.strip()
        start, end = resp_text.find("{"), resp_text.rfind("}")
        if start != -1 and end != -1:
            resp_text = resp_text[start:end + 1]
        result = json.loads(resp_text)
        result["module_id"] = module_id
        result["assessed_at"] = datetime.now(timezone.utc).isoformat()
        await db.assessments.update_one({"module_id": module_id}, {"$set": result}, upsert=True)
        return {"status": "success", "assessment": result}
    except Exception as e:
        logger.error(f"Assessment error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/modules/{module_id}/assessment")
async def get_assessment(module_id: str):
    a = await db.assessments.find_one({"module_id": module_id}, {"_id": 0})
    if not a:
        raise HTTPException(status_code=404, detail="No assessment found")
    return a

# ─── ROLE PLAY REVIEW ENDPOINT ───────────────────────────────────────────

@api_router.post("/modules/{module_id}/roleplay")
async def review_role_play(module_id: str, file: UploadFile = File(...), request: Request = None):
    module = await db.modules.find_one({"id": module_id}, {"_id": 0})
    if not module:
        raise HTTPException(status_code=404, detail="Module not found")

    text = await read_uploaded_file(file)

    channel = module.get("channel", "AMZ")

    try:
        from openai import AsyncOpenAI as _OpenAI
        api_key = os.environ.get("EMERGENT_LLM_KEY")
        system_prompt = f"""You are a role play evaluator for Quartile Academy. Evaluate this {channel} role play / QBR presentation. Return ONLY valid JSON:
{{"student_name":"Name","overall_score":85,"channel_knowledge":{{"score":8,"feedback":"How well they demonstrated {channel} platform knowledge"}},"data_clarity":{{"score":7,"feedback":"detail"}},"storytelling":{{"score":8,"feedback":"detail"}},"analytical_thinking":{{"score":7,"feedback":"detail"}},"communication":{{"score":9,"feedback":"detail"}},"soft_skills":{{"score":8,"feedback":"detail"}},"presentation_design":{{"score":7,"feedback":"detail"}},"optimization_opportunities":{{"score":8,"feedback":"detail"}},"upsell_opportunities":{{"score":7,"feedback":"detail"}},"negotiation":{{"score":8,"feedback":"detail"}},"objection_handling":{{"score":7,"feedback":"detail"}},"asking_questions":{{"score":8,"feedback":"detail"}},"presentation_pacing":{{"score":7,"feedback":"detail"}},"presentation_time":{{"score":8,"feedback":"Appropriate length and time management"}},"overall_feedback":"2-3 sentence summary","strengths":["s1","s2"],"areas_for_improvement":["a1","a2"]}}
Score each dimension 1-10. Be thorough."""

        _client = _OpenAI(api_key=api_key)
        user_text = f"Role Play for {module.get('module', '')} ({channel}):\n\n{text[:15000]}"
        _completion = await _client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_text}
            ]
        )
        resp = _completion.choices[0].message.content
        resp_text = resp.strip()
        start, end = resp_text.find("{"), resp_text.rfind("}")
        if start != -1 and end != -1:
            resp_text = resp_text[start:end + 1]
        result = json.loads(resp_text)
        result["module_id"] = module_id
        result["channel"] = channel
        result["reviewed_at"] = datetime.now(timezone.utc).isoformat()
        await db.roleplay_reviews.update_one({"module_id": module_id}, {"$set": result}, upsert=True)
        return {"status": "success", "review": result}
    except Exception as e:
        logger.error(f"Role play review error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/modules/{module_id}/roleplay")
async def get_role_play_review(module_id: str):
    r = await db.roleplay_reviews.find_one({"module_id": module_id}, {"_id": 0})
    if not r:
        raise HTTPException(status_code=404, detail="No role play review found")
    return r

# ─── API ROUTES ──────────────────────────────────────────────────────────────

@api_router.get("/")
async def root():
    return {"message": "Quartile Academic Tracker API"}

# Dashboard
@api_router.get("/dashboard")
async def get_dashboard():
    modules = await db.modules.find({}, {"_id": 0}).to_list(1000)
    presented = [m for m in modules if m.get("status") == "Presented"]
    total_hrs = sum(m.get("length_hrs", 0) for m in modules)
    presented_hrs = sum(m.get("length_hrs", 0) for m in presented)
    instructors = await db.instructors.count_documents({})
    analyzed = await db.modules.count_documents({"analyzed": True})

    weeks = sorted(set(m["week"] for m in modules))
    weekly_progress = []
    for w in weeks:
        week_mods = [m for m in modules if m["week"] == w]
        pres = len([m for m in week_mods if m["status"] == "Presented"])
        weekly_progress.append({"week": w, "total": len(week_mods), "presented": pres, "pct": round(pres / len(week_mods) * 100) if week_mods else 0})

    upcoming = [m for m in modules if m.get("status") == "Scheduled"][:5]

    return {
        "modules_delivered": len(presented),
        "total_hrs": round(total_hrs, 1),
        "presented_hrs": round(presented_hrs, 1),
        "active_instructors": instructors,
        "modules_analyzed": analyzed,
        "total_modules": len(modules),
        "weekly_progress": weekly_progress,
        "upcoming_modules": upcoming,
    }

# Modules
@api_router.get("/modules")
async def get_modules(week: Optional[str] = None, search: Optional[str] = None, channel: Optional[str] = None):
    query = {}
    if week:
        query["week"] = week
    if channel:
        query["channel"] = {"$regex": channel, "$options": "i"}
    if search:
        query["$or"] = [
            {"module": {"$regex": search, "$options": "i"}},
            {"instructor": {"$regex": search, "$options": "i"}},
            {"id": {"$regex": search, "$options": "i"}},
        ]
    modules = await db.modules.find(query, {"_id": 0}).to_list(1000)
    return modules

@api_router.get("/modules/{module_id}")
async def get_module(module_id: str):
    module = await db.modules.find_one({"id": module_id}, {"_id": 0})
    if not module:
        raise HTTPException(status_code=404, detail="Module not found")
    analysis = await db.analyses.find_one({"module_id": module_id}, {"_id": 0})
    module["analysis"] = analysis
    content_doc = await db.module_content.find_one({"module_id": module_id}, {"_id": 0, "raw_text": 0})
    module["has_content"] = content_doc is not None
    if content_doc:
        module["content"] = content_doc
    # Assessment data
    assessment = await db.assessments.find_one({"module_id": module_id}, {"_id": 0})
    if assessment:
        module["assessment"] = assessment
    # Role play data
    roleplay = await db.roleplay_reviews.find_one({"module_id": module_id}, {"_id": 0})
    if roleplay:
        module["roleplay"] = roleplay
    return module

# ─── Module Content Upload ────────────────────────────────────────────────

@api_router.post("/modules/{module_id}/content")
async def upload_module_content(module_id: str, file: UploadFile = File(...)):
    module = await db.modules.find_one({"id": module_id}, {"_id": 0})
    if not module:
        raise HTTPException(status_code=404, detail="Module not found")

    fname = (file.filename or "").lower()
    if not (fname.endswith(".docx") or fname.endswith(".txt")):
        raise HTTPException(status_code=400, detail="Only .docx and .txt files are supported")

    file_content = await file.read()

    sections = []
    current_section = {"title": "Introduction", "content": []}
    topics = []
    learning_objectives = []

    if fname.endswith(".txt"):
        raw_text = file_content.decode("utf-8", errors="replace")
        for line in raw_text.split("\n"):
            text = line.strip()
            if not text:
                continue
            if any(text.startswith(p) for p in ["I.", "II.", "III.", "IV.", "V.", "VI.", "VII.", "VIII."]):
                if current_section["content"] or current_section["title"] != "Introduction":
                    sections.append(current_section)
                current_section = {"title": text, "content": []}
            elif text.startswith(("A.", "B.", "C.", "D.", "E.", "F.", "G.")):
                topic_name = text.lstrip("ABCDEFG. ").strip()
                if topic_name:
                    topics.append(topic_name)
                current_section["content"].append(text)
            elif "?" in text and len(text) > 20:
                learning_objectives.append(text)
                current_section["content"].append(text)
            else:
                current_section["content"].append(text)
        if current_section["content"] or current_section["title"] != "Introduction":
            sections.append(current_section)
    else:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name
        try:
            import docx
            doc = docx.Document(tmp_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = (para.style.name or "").lower()
                is_heading = "heading" in style_name or "title" in style_name

                if is_heading or any(text.startswith(p) for p in ["I.", "II.", "III.", "IV.", "V.", "VI.", "VII.", "VIII."]):
                    if current_section["content"] or current_section["title"] != "Introduction":
                        sections.append(current_section)
                    current_section = {"title": text, "content": []}
                elif text.startswith(("A.", "B.", "C.", "D.", "E.", "F.", "G.")):
                    topic_name = text.lstrip("ABCDEFG. ").strip()
                    if topic_name:
                        topics.append(topic_name)
                    current_section["content"].append(text)
                elif "?" in text and len(text) > 20:
                    learning_objectives.append(text)
                    current_section["content"].append(text)
                elif para.runs and any(run.bold for run in para.runs) and len(text) < 120:
                    bold_text = "".join([r.text for r in para.runs if r.bold]).strip()
                    if bold_text and len(bold_text) > 5:
                        topics.append(bold_text)
                    current_section["content"].append(text)
                else:
                    current_section["content"].append(text)

            if current_section["content"] or current_section["title"] != "Introduction":
                sections.append(current_section)
        finally:
            os.unlink(tmp_path)

    if not topics:
        topics = [s["title"] for s in sections if s["title"] != "Introduction" and len(s["title"]) < 120]

    if len(raw_text) < 50:
        raise HTTPException(status_code=400, detail="Document appears empty or too short")

    content_doc = {
        "module_id": module_id,
        "filename": file.filename,
        "raw_text": raw_text,
        "sections": [{"title": s["title"], "content": "\n".join(s["content"])} for s in sections],
        "topics": topics if topics else [s["title"] for s in sections if s["title"] != "Introduction"],
        "learning_objectives": learning_objectives,
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    }

    await db.module_content.update_one(
        {"module_id": module_id},
        {"$set": content_doc},
        upsert=True
    )

    # Update module flag
    await db.modules.update_one({"id": module_id}, {"$set": {"has_content": True}})

    return {"status": "success", "sections": len(sections), "topics": content_doc["topics"], "learning_objectives": content_doc["learning_objectives"]}

@api_router.get("/modules/{module_id}/content")
async def get_module_content(module_id: str):
    content_doc = await db.module_content.find_one({"module_id": module_id}, {"_id": 0})
    if not content_doc:
        raise HTTPException(status_code=404, detail="No module content uploaded")
    return content_doc

# Transcript Ingestion
@api_router.post("/modules/{module_id}/ingest")
async def ingest_transcript(module_id: str, file: UploadFile = File(...)):
    module = await db.modules.find_one({"id": module_id}, {"_id": 0})
    if not module:
        raise HTTPException(status_code=404, detail="Module not found")

    fname = (file.filename or "").lower()
    if not (fname.endswith(".docx") or fname.endswith(".txt")):
        raise HTTPException(status_code=400, detail="Only .docx and .txt files are supported")

    content = await file.read()

    if fname.endswith(".txt"):
        transcript_text = content.decode("utf-8", errors="replace")
    else:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            import docx
            doc = docx.Document(tmp_path)
            transcript_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        finally:
            os.unlink(tmp_path)

    if len(transcript_text) < 100:
        raise HTTPException(status_code=400, detail="Transcript appears to be too short or empty")

    # AI Analysis
    try:
        from openai import AsyncOpenAI as _OpenAI
        api_key = os.environ.get("EMERGENT_LLM_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="LLM API key not configured")

        # Check if module content exists for reference-based analysis
        module_content = await db.module_content.find_one({"module_id": module_id}, {"_id": 0, "raw_text": 1, "topics": 1, "learning_objectives": 1})

        if module_content:
            # Enhanced prompt: match transcript AGAINST the module content
            content_ref = module_content.get("raw_text", "")[:5000]
            content_topics = ", ".join(module_content.get("topics", []))
            content_objectives = "\n".join(module_content.get("learning_objectives", []))

            system_prompt = f"""You are an academic content analyst for Quartile Academy. You are given a MODULE CONTENT DOCUMENT (the source of truth) and a SESSION TRANSCRIPT. Your job is to compare the transcript against the module content to evaluate how well the instructor covered the required material.

MODULE CONTENT REFERENCE (source of truth topics):
Topics: {content_topics}
Learning Objectives: {content_objectives}
Full content excerpt: {content_ref}

Return ONLY valid JSON (no markdown, no ```). Use this exact structure:
{{"learning_objective":"1-2 sentence summary based on the MODULE CONTENT","content_match_score":9,"content_match_summary":"3-4 sentence assessment comparing transcript vs module content. What % was covered, what was missed, what was added beyond scope","topics_covered":["topic1","topic2"],"topics_missed":["topics from MODULE CONTENT that were NOT covered in transcript"],"terminology_drifts":[{{"id":"Correction #01","severity":"CRITICAL","issue":"desc","recommendation":"fix"}}],"tips_shared":["tip1","tip2"],"follow_up_items":["item1"],"qa_log":[{{"student":"Name","question":"q","answer":"a","section":"Section"}}],"student_performance":[{{"name":"Name","interactions":10,"score":8,"strengths":"desc","areas_for_improvement":"desc","skills":{{"objection_handling":7,"negotiation":6,"data_analysis":8,"communication":9,"presentation":7,"analytical_thinking":8,"campaign_management":7,"client_management":6}}}}],"overall_score":9,"avg_satisfaction":4.5,"session_sections":[{{"title":"Title","start_time":"09:00","end_time":"09:15","summary":"Brief"}}],"instructor_name":"Full Name","actual_duration":"2h 05m","participants":["Name1","Name2"]}}
CRITICAL: For content_match_score, score STRICTLY based on how well the transcript covers the MODULE CONTENT topics. topics_missed MUST list topics from the module content that were NOT adequately covered. For each student, rate ALL 8 skills 0-10. Be thorough."""
        else:
            system_prompt = """You are an academic content analyst for Quartile Academy. Analyze the transcript and return ONLY valid JSON (no markdown, no ```). Use this exact structure:
{"learning_objective":"1-2 sentence summary","content_match_score":9,"content_match_summary":"3-4 sentence assessment","topics_covered":["topic1","topic2"],"topics_missed":["topic1"],"terminology_drifts":[{"id":"Correction #01","severity":"CRITICAL","issue":"desc","recommendation":"fix"}],"tips_shared":["tip1","tip2"],"follow_up_items":["item1"],"qa_log":[{"student":"Name","question":"q","answer":"a","section":"Section"}],"student_performance":[{"name":"Name","interactions":10,"score":8,"strengths":"desc","areas_for_improvement":"desc","skills":{"objection_handling":7,"negotiation":6,"data_analysis":8,"communication":9,"presentation":7,"analytical_thinking":8,"campaign_management":7,"client_management":6}}],"overall_score":9,"avg_satisfaction":4.5,"session_sections":[{"title":"Title","start_time":"09:00","end_time":"09:15","summary":"Brief"}],"instructor_name":"Full Name","actual_duration":"2h 05m","participants":["Name1","Name2"]}
For each student in student_performance, rate ALL 8 skills from 0-10 based on their demonstrated ability in the transcript. If a skill wasn't demonstrated, estimate based on their overall engagement. Extract ALL Q&A interactions, tips, student performance. Be thorough."""

        # Trim transcript to ~15k chars to reduce token cost
        trimmed = transcript_text[:15000]
        if len(transcript_text) > 15000:
            # Also grab the last 3k chars for closing sections
            trimmed += "\n...[MIDDLE SECTIONS TRIMMED]...\n" + transcript_text[-3000:]

        _client = _OpenAI(api_key=api_key)
        user_text = f"Analyze this module transcript for module {module_id} - {module.get('module', '')}:\n\n{trimmed}"
        _completion = await _client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_text}
            ]
        )
        response = _completion.choices[0].message.content

        # Parse JSON from response
        response_text = response.strip()
        # Remove markdown code fences if present
        if "```" in response_text:
            # Find JSON between code fences
            import re
            json_match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', response_text, re.DOTALL)
            if json_match:
                response_text = json_match.group(1).strip()
            else:
                # Fallback: strip leading/trailing ```
                response_text = response_text.replace("```json", "").replace("```", "").strip()
        
        # Find the first { and last } to extract JSON object
        start = response_text.find("{")
        end = response_text.rfind("}")
        if start != -1 and end != -1:
            response_text = response_text[start:end + 1]

        analysis = json.loads(response_text.strip())
        analysis["module_id"] = module_id
        analysis["ingested_at"] = datetime.now(timezone.utc).isoformat()
        # ALWAYS use system instructor name over AI-extracted name
        analysis["instructor_name"] = module.get("instructor", analysis.get("instructor_name", ""))
        # Flag if module content was used for matching
        analysis["content_matched"] = module_content is not None

        # Upsert analysis
        await db.analyses.update_one(
            {"module_id": module_id},
            {"$set": analysis},
            upsert=True
        )

        # Update module status
        await db.modules.update_one(
            {"id": module_id},
            {"$set": {"analyzed": True, "status": "Presented"}}
        )

        # ─── Update student records from analysis ─────────────────────────
        await update_student_records(module_id, module, analysis)

        return {"status": "success", "analysis": analysis}

    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse AI response: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse AI analysis response")
    except Exception as e:
        logger.error(f"Transcript ingestion error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ─── STUDENT RECORD UPDATER (called after transcript ingestion) ───────────

async def update_student_records(module_id: str, module: dict, analysis: dict):
    """Cross-reference analysis student_performance with module metadata to update student KPIs."""
    student_perfs = analysis.get("student_performance", [])
    if not student_perfs:
        return

    channel = module.get("channel", "")
    hrs = module.get("length_hrs", 0)
    mod_name = module.get("module", "")
    mod_date = module.get("date", "")
    is_shadow = "shadow" in mod_name.lower()
    is_buddy = module.get("instructor", "").strip().lower() == "buddy" or "buddy" in mod_name.lower()
    is_assessment = module.get("format", "").lower() == "assessment"

    for sp in student_perfs:
        student_name = sp.get("name", "")
        if not student_name:
            continue

        # Find student by fuzzy name match
        student = await db.students.find_one(
            {"name": {"$regex": f"^{student_name[:4]}", "$options": "i"}},
            {"_id": 1, "modules_attended": 1, "channel_hours": 1, "skills": 1,
             "shadow_calls": 1, "buddy_tasks": 1, "assessment_scores": 1}
        )
        if not student:
            continue

        sid = student["_id"]
        existing_mods = [m["module_id"] for m in student.get("modules_attended", [])]
        if module_id in existing_mods:
            continue  # Already recorded

        attendance_record = {
            "module_id": module_id,
            "module": mod_name,
            "date": mod_date,
            "score": sp.get("score", 0),
            "channel": channel,
            "hrs": hrs,
            "interactions": sp.get("interactions", 0),
        }

        # Channel hours update
        ch_key = f"channel_hours.{channel}"
        current_ch_hrs = student.get("channel_hours", {}).get(channel, 0)
        new_ch_hrs = current_ch_hrs + hrs

        update_ops = {
            "$push": {"modules_attended": attendance_record},
            "$inc": {
                "sessions": 1,
                "total_hours": hrs,
            },
            "$set": {ch_key: round(new_ch_hrs, 1)},
        }

        if is_shadow:
            update_ops["$inc"]["shadow_calls"] = 1
        if is_buddy:
            update_ops["$inc"]["buddy_tasks"] = 1

        # Update skills (rolling average)
        sp_skills = sp.get("skills", {})
        if sp_skills:
            existing_skills = student.get("skills", {})
            n_mods = len(existing_mods) + 1
            for skill_key in ["objection_handling", "negotiation", "data_analysis", "communication",
                              "presentation", "analytical_thinking", "campaign_management", "client_management"]:
                new_val = sp_skills.get(skill_key, 0)
                if new_val > 0:
                    old_val = existing_skills.get(skill_key, 0)
                    # Weighted moving average: give more weight to recent scores
                    avg = round(((old_val * (n_mods - 1)) + new_val) / n_mods, 1)
                    update_ops["$set"][f"skills.{skill_key}"] = avg

        # Assessment score tracking
        if is_assessment:
            update_ops.setdefault("$push", {})
            update_ops["$push"]["assessment_scores"] = {
                "module_id": module_id,
                "title": mod_name,
                "score": sp.get("score", 0) * 10,  # Convert to 0-100
            }

        # Strengths and improvement areas
        if sp.get("strengths"):
            update_ops["$set"]["strengths_summary"] = sp["strengths"]
        if sp.get("areas_for_improvement"):
            update_ops["$set"]["improvement_areas"] = sp["areas_for_improvement"]

        await db.students.update_one({"_id": sid}, update_ops)

        # Recompute academic_progress and overall_score
        updated = await db.students.find_one({"_id": sid}, {"_id": 0, "modules_attended": 1})
        mods_att = updated.get("modules_attended", [])
        if mods_att:
            total_presented = await db.modules.count_documents({"status": "Presented"})
            attendance_pct = round(len(mods_att) / max(total_presented, 1) * 100)
            avg_score = round(sum(m.get("score", 0) for m in mods_att) / len(mods_att), 1)
            await db.students.update_one({"_id": sid}, {
                "$set": {"academic_progress": min(attendance_pct, 100), "overall_score": avg_score}
            })

    logger.info(f"Updated {len(student_perfs)} student records for module {module_id}")

# ─── STUDENT DETAIL ENDPOINT ─────────────────────────────────────────────

@api_router.get("/students/{student_id}")
async def get_student_detail(student_id: str):
    student = await db.students.find_one({"student_id": student_id}, {"_id": 0})
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # Compute derived KPIs
    mods = student.get("modules_attended", [])
    total_presented = await db.modules.count_documents({"status": "Presented"})
    student["attendance_rate"] = round(len(mods) / max(total_presented, 1) * 100) if mods else 0

    # Per-module Q&A participation (from all analyses)
    qa_entries = []
    for m in mods:
        analysis = await db.analyses.find_one({"module_id": m["module_id"]}, {"_id": 0, "qa_log": 1})
        if analysis and analysis.get("qa_log"):
            for qa in analysis["qa_log"]:
                if student["name"].split()[0].lower() in qa.get("student", "").lower():
                    qa_entries.append({**qa, "module": m["module"], "module_id": m["module_id"]})
    student["qa_history"] = qa_entries

    return student

# Instructors
@api_router.get("/instructors")
async def get_instructors(channel: Optional[str] = None):
    query = {}
    if channel:
        query["channels"] = {"$regex": channel, "$options": "i"}
    instructors = await db.instructors.find(query, {"_id": 0}).to_list(1000)
    return instructors

# Students
@api_router.get("/students")
async def get_students(search: Optional[str] = None):
    query = {}
    if search:
        query["name"] = {"$regex": search, "$options": "i"}
    students = await db.students.find(query, {"_id": 0}).to_list(1000)
    return students

# Calendar data
@api_router.get("/calendar")
async def get_calendar(week: Optional[str] = None):
    query = {}
    if week:
        query["week"] = week
    modules = await db.modules.find(query, {"_id": 0}).to_list(1000)
    return modules

# Unique filter values
@api_router.get("/filters")
async def get_filters():
    modules = await db.modules.find({}, {"_id": 0, "week": 1, "channel": 1}).to_list(1000)
    weeks = sorted(set(m["week"] for m in modules))
    channels = sorted(set(m["channel"] for m in modules))
    return {"weeks": weeks, "channels": channels}

app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)
